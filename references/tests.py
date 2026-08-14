import socket
from io import BytesIO
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from .forms import ReferenceForm
from .models import Project, ProjectReference, Reference
from .services.citations import generate_citations
from .services.extraction import extract_html
from .services.people import parse_people
from .services.security import UnsafeURL, validate_public_url


class BaseDataTest(TestCase):
    def setUp(self):
        User = get_user_model()
        self.user = User.objects.create_user(
            "rodrigo", "rodrigo@example.com", "SenhaSegura123!"
        )
        self.other = User.objects.create_user(
            "outro", "outro@example.com", "SenhaSegura123!"
        )
        self.project = Project.objects.create(
            owner=self.user, name="Monografia", default_style="abnt"
        )
        self.reference = Reference.objects.create(
            owner=self.user,
            reference_type=Reference.Type.JOURNAL_ARTICLE,
            authors=parse_people("Silva, João Paulo; Souza, Maria"),
            title="Inovação e cuidado em saúde",
            year="2025",
            container_title="Revista Brasileira de Saúde",
            volume="12",
            issue="2",
            pages="15-28",
            doi="10.1234/exemplo.2025",
            extraction_status=Reference.ExtractionStatus.REVIEWED,
        )
        ProjectReference.objects.create(
            project=self.project, reference=self.reference, position=1
        )


class ModelAndCitationTests(BaseDataTest):
    def test_doi_is_normalized_and_fingerprint_created(self):
        reference = Reference.objects.create(
            owner=self.user,
            title="Outro artigo",
            doi="https://doi.org/10.9999/Teste",
        )
        self.assertEqual(reference.doi, "10.9999/Teste")
        self.assertEqual(len(reference.normalized_fingerprint), 64)

    def test_abnt_contains_expected_elements(self):
        text = generate_citations([self.reference], "abnt")[0].text
        self.assertIn("SILVA, João Paulo", text)
        self.assertIn("v. 12", text)
        self.assertIn("DOI: https://doi.org/10.1234/exemplo.2025", text)

    def test_apa_and_numeric_styles(self):
        apa = generate_citations([self.reference], "apa")[0].text
        self.assertIn("Silva, J. P.", apa)
        self.assertIn("(2025)", apa)
        vancouver = generate_citations([self.reference], "vancouver")[0].text
        ieee = generate_citations([self.reference], "ieee")[0].text
        self.assertTrue(vancouver.startswith("1."))
        self.assertTrue(ieee.startswith("[1]"))

    def test_reference_form_parses_authors(self):
        data = {
            "reference_type": Reference.Type.BOOK,
            "title": "Livro de teste",
            "authors_input": "Niskier, Rodrigo\nInstituição: Organização Mundial da Saúde",
            "editors_input": "",
            "year": "2026",
            "language": "pt",
        }
        form = ReferenceForm(data=data)
        self.assertTrue(form.is_valid(), form.errors)
        instance = form.save(commit=False)
        self.assertEqual(instance.authors[0]["family"], "Niskier")
        self.assertEqual(instance.authors[1]["literal"], "Organização Mundial da Saúde")


class ExtractionTests(TestCase):
    @override_settings(CROSSREF_ENRICH=False)
    def test_extracts_academic_meta_tags_from_html(self):
        html = b"""
        <html><head>
        <meta name="citation_title" content="A Reliable Study">
        <meta name="citation_author" content="Jane Doe">
        <meta name="citation_publication_date" content="2024-05-12">
        <meta name="citation_journal_title" content="Evidence Journal">
        <meta name="citation_doi" content="10.5555/reliable">
        </head></html>
        """
        result = extract_html(html, "https://example.org/article")
        self.assertEqual(result.fields["title"], "A Reliable Study")
        self.assertEqual(result.fields["year"], "2024")
        self.assertEqual(result.fields["doi"], "10.5555/reliable")
        self.assertEqual(result.fields["authors"][0]["family"], "Doe")

    @patch("references.services.security.socket.getaddrinfo")
    def test_ssrf_blocks_private_addresses(self, mocked_dns):
        mocked_dns.return_value = [
            (socket.AF_INET, socket.SOCK_STREAM, 6, "", ("127.0.0.1", 80))
        ]
        with self.assertRaises(UnsafeURL):
            validate_public_url("http://example.org/internal")


class ViewPermissionTests(BaseDataTest):
    def setUp(self):
        super().setUp()
        self.client.login(username="rodrigo", password="SenhaSegura123!")

    def test_user_cannot_open_another_users_project(self):
        foreign = Project.objects.create(owner=self.other, name="Privado")
        response = self.client.get(
            reverse("references:project_detail", args=[foreign.pk])
        )
        self.assertEqual(response.status_code, 404)

    def test_citarn_uses_shared_login_and_namespaced_assets(self):
        response = self.client.get(reverse("references:dashboard"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "/static/references/css/app.css")
        self.assertContains(response, reverse("works:dashboard"))

        self.client.logout()
        protected = self.client.get(reverse("references:dashboard"))
        self.assertRedirects(
            protected,
            f"{reverse('login')}?next={reverse('references:dashboard')}",
        )

    def test_user_cannot_download_another_users_file(self):
        foreign = Reference.objects.create(
            owner=self.other,
            title="Arquivo privado",
            source_file=SimpleUploadedFile("privado.pdf", b"%PDF-1.4 fake"),
        )
        response = self.client.get(
            reverse("references:reference_file", args=[foreign.pk])
        )
        self.assertEqual(response.status_code, 404)

    def test_project_generation_saves_selection_and_style(self):
        response = self.client.post(
            reverse("references:generate_list", args=[self.project.pk]),
            {"style": "apa", "references": [str(self.reference.pk)]},
        )
        self.assertRedirects(
            response, reverse("references:generate_list", args=[self.project.pk])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.default_style, "apa")

    def test_export_docx_pdf_and_txt(self):
        for extension, expected_type in [
            (
                "docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            ("pdf", "application/pdf"),
            ("txt", "text/plain"),
        ]:
            response = self.client.get(
                reverse("references:export_project", args=[self.project.pk, extension])
            )
            self.assertEqual(response.status_code, 200)
            self.assertTrue(response["Content-Type"].startswith(expected_type))
            self.assertGreater(len(response.content), 50)

    @override_settings(CROSSREF_ENRICH=False)
    def test_imports_docx_and_links_it_to_project(self):
        document = Document()
        document.core_properties.title = "Documento científico de teste"
        document.core_properties.author = "Rodrigo Niskier"
        document.add_paragraph("Publicado em 2026. DOI: 10.7777/teste")
        output = BytesIO()
        document.save(output)
        upload = SimpleUploadedFile(
            "artigo.docx",
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = self.client.post(
            reverse("references:import_sources", args=[self.project.pk]),
            {"files": [upload], "urls": ""},
        )
        self.assertRedirects(
            response, reverse("references:project_detail", args=[self.project.pk])
        )
        imported = Reference.objects.get(owner=self.user, doi="10.7777/teste")
        self.assertEqual(imported.title, "Documento científico de teste")
        self.assertTrue(
            ProjectReference.objects.filter(
                project=self.project, reference=imported
            ).exists()
        )

    def test_remove_from_project_keeps_library_reference(self):
        response = self.client.post(
            reverse(
                "references:remove_from_project",
                args=[self.project.pk, self.reference.pk],
            )
        )
        self.assertRedirects(response, self.project.get_absolute_url())
        self.assertTrue(Reference.objects.filter(pk=self.reference.pk).exists())
        self.assertFalse(
            ProjectReference.objects.filter(
                project=self.project, reference=self.reference
            ).exists()
        )


class FullStoryTests(TestCase):
    def test_shared_account_to_export_story(self):
        User = get_user_model()
        user = User.objects.create_user(
            username="ana.pesquisa",
            email="ana@example.com",
            password="UmaSenhaAcademica123!",
            first_name="Ana",
            last_name="Pesquisadora",
        )
        self.client.force_login(user)

        dashboard = self.client.get(reverse("references:dashboard"))
        self.assertEqual(dashboard.status_code, 200)
        self.assertContains(dashboard, "Ana")

        project_response = self.client.post(
            reverse("references:project_create"),
            {
                "name": "Artigo de revisão",
                "description": "Fluxo integral",
                "default_style": "abnt",
            },
        )
        project = Project.objects.get(name="Artigo de revisão")
        self.assertRedirects(
            project_response, reverse("references:import_sources", args=[project.pk])
        )

        reference_response = self.client.post(
            reverse("references:reference_create"),
            {
                "project_id": str(project.pk),
                "reference_type": Reference.Type.BOOK,
                "title": "Metodologia da pesquisa clínica",
                "authors_input": "Costa, Ana Maria",
                "editors_input": "",
                "year": "2026",
                "publisher": "Editora Acadêmica",
                "publisher_place": "São Paulo",
                "language": "pt",
            },
        )
        reference = Reference.objects.get(title="Metodologia da pesquisa clínica")
        self.assertRedirects(reference_response, project.get_absolute_url())
        self.assertTrue(
            ProjectReference.objects.filter(
                project=project, reference=reference
            ).exists()
        )

        generate_response = self.client.post(
            reverse("references:generate_list", args=[project.pk]),
            {
                "style": "apa",
                "references": [str(reference.pk)],
                f"position_{reference.pk}": "1",
            },
        )
        self.assertRedirects(
            generate_response, reverse("references:generate_list", args=[project.pk])
        )
        preview = self.client.get(
            reverse("references:generate_list", args=[project.pk])
        )
        self.assertContains(preview, "Metodologia da pesquisa clínica")

        export_response = self.client.get(
            reverse("references:export_project", args=[project.pk, "docx"])
        )
        self.assertEqual(export_response.status_code, 200)
        exported = Document(BytesIO(export_response.content))
        exported_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        self.assertIn("REFERÊNCIAS", exported_text)
        self.assertIn("Metodologia da pesquisa clínica", exported_text)
