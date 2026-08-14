from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate

from .citations import generate_citations, split_emphasis


def export_txt(references, style: str) -> bytes:
    items = generate_citations(references, style)
    body = "REFERÊNCIAS\n\n" + "\n\n".join(item.text for item in items)
    return body.encode("utf-8")


def export_docx(references, style: str, project_name: str) -> bytes:
    items = generate_citations(references, style)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("REFERÊNCIAS")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)

    for item in items:
        paragraph = document.add_paragraph()
        for value, emphasized in split_emphasis(item.text, item.reference):
            run = paragraph.add_run(value)
            run.italic = emphasized
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style == "apa":
            paragraph.paragraph_format.line_spacing = 2
            paragraph.paragraph_format.left_indent = Cm(1.27)
            paragraph.paragraph_format.first_line_indent = Cm(-1.27)
        else:
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_after = Pt(12)
        if style in {"chicago", "harvard"}:
            paragraph.paragraph_format.left_indent = Cm(1.27)
            paragraph.paragraph_format.first_line_indent = Cm(-1.27)

    document.core_properties.title = f"Referências — {project_name}"
    document.core_properties.subject = f"Lista de referências no estilo {style.upper()}"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_footer(canvas, document):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColorRGB(0.35, 0.39, 0.45)
    canvas.drawCentredString(A4[0] / 2, 1.2 * cm, str(document.page))
    canvas.restoreState()


def export_pdf(references, style: str, project_name: str) -> bytes:
    items = generate_citations(references, style)
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=3 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.2 * cm,
        title=f"Referências — {project_name}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReferenceTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "ReferenceBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=22 if style == "apa" else 14,
        alignment=TA_LEFT,
        spaceAfter=12,
        leftIndent=1.27 * cm if style in {"apa", "chicago", "harvard"} else 0,
        firstLineIndent=-1.27 * cm if style in {"apa", "chicago", "harvard"} else 0,
    )
    story = [Paragraph("REFERÊNCIAS", title_style)]
    if not items:
        story.append(Paragraph("Nenhuma referência selecionada.", body_style))
    for item in items:
        markup = "".join(
            f"<i>{escape(value)}</i>" if emphasized else escape(value)
            for value, emphasized in split_emphasis(item.text, item.reference)
        )
        story.append(Paragraph(markup, body_style))
    document.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    return output.getvalue()
