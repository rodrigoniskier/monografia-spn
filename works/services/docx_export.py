"""Geração determinística da monografia SPN em DOCX segundo a ABNT atual."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import Iterable
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .abnt import format_reference
from .docx_footnotes import inject_footnotes


FONT_NAME = "Times New Roman"
BODY_SIZE = 12
SMALL_SIZE = 10
SPN_LOGO_PATH = (
    Path(__file__).resolve().parents[1]
    / "static"
    / "works"
    / "img"
    / "spn-logo.png"
)
def _safe(value) -> str:
    text = str(value or "")
    return "".join(char for char in text if char in "\n\t" or ord(char) >= 32).strip()


def _set_run_font(run, *, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.7)
    section.footer_distance = Cm(1.2)


def _set_page_number_start(section, start: int):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is None:
        existing = OxmlElement("w:pgNumType")
        sect_pr.append(existing)
    existing.set(qn("w:start"), str(start))


def _continue_page_numbering(section):
    """Remove o reinício copiado pelo python-docx ao criar uma nova seção."""
    existing = section._sectPr.find(qn("w:pgNumType"))
    if existing is not None:
        section._sectPr.remove(existing)


def _add_page_number(section):
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    _set_run_font(run, size=10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def _configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for level in range(1, 6):
        name = f"Heading {level}"
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(BODY_SIZE)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True

    if "Long Quotation" not in [style.name for style in doc.styles]:
        quote = doc.styles.add_style("Long Quotation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = doc.styles["Long Quotation"]
    quote.font.name = FONT_NAME
    quote._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    quote._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    quote.font.size = Pt(SMALL_SIZE)
    quote.paragraph_format.left_indent = Cm(4)
    quote.paragraph_format.first_line_indent = Cm(0)
    quote.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    quote.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(6)

    for list_name, left, first in (("List Bullet", 1.25, -0.63), ("List Number", 1.25, -0.63)):
        style = doc.styles[list_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(BODY_SIZE)
        style.paragraph_format.left_indent = Cm(left)
        style.paragraph_format.first_line_indent = Cm(first)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    if "Footnote Text" not in [style.name for style in doc.styles]:
        footnote_text = doc.styles.add_style("Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        footnote_text = doc.styles["Footnote Text"]
    footnote_text.font.name = FONT_NAME
    footnote_text._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    footnote_text._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    footnote_text.font.size = Pt(SMALL_SIZE)
    footnote_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    footnote_text.paragraph_format.first_line_indent = Cm(0)
    footnote_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    footnote_text.paragraph_format.space_before = Pt(0)
    footnote_text.paragraph_format.space_after = Pt(0)

    if "Footnote Reference" not in [style.name for style in doc.styles]:
        footnote_reference = doc.styles.add_style(
            "Footnote Reference", WD_STYLE_TYPE.CHARACTER
        )
    else:
        footnote_reference = doc.styles["Footnote Reference"]
    footnote_reference.font.name = FONT_NAME
    footnote_reference._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    footnote_reference._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    footnote_reference.font.size = Pt(SMALL_SIZE)
    footnote_reference.font.superscript = True


def _add_centered(doc, text, *, bold=False, size=BODY_SIZE, after=0, before=0, uppercase=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(_safe(text).upper() if uppercase else _safe(text))
    _set_run_font(run, size=size, bold=bold)
    return paragraph


def _add_spacer(doc, points: float):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return paragraph


def _add_spn_logo(doc):
    """Insere a marca usada nas capas recentes do SPN."""
    if not SPN_LOGO_PATH.is_file():
        return None
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(SPN_LOGO_PATH), width=Cm(5.7))
    return paragraph


def _bottom_block_spacer(text: str, *, base=590, reduction_per_line=18) -> float:
    """Aproxima o bloco da borda inferior sem estourar textos mais longos."""
    raw_lines = _safe(text).splitlines() or [""]
    estimated_lines = sum(max(1, (len(line) + 42) // 43) for line in raw_lines)
    return max(300, base - max(estimated_lines - 2, 0) * reduction_per_line)


def _page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _pretext_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.line_spacing = 1.5
    _set_run_font(paragraph.add_run(_safe(text).upper()), bold=True)
    return paragraph


def _add_prose(doc, text: str):
    """Converte texto simples em parágrafos, listas e citações longas marcadas com >."""
    text = _safe(text)
    if not text:
        return
    blocks = re.split(r"\n\s*\n", text)
    for block in blocks:
        lines = [line.rstrip() for line in block.splitlines()]
        if all(not line.strip() for line in lines):
            continue
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith(">"):
                paragraph = doc.add_paragraph(style="Long Quotation")
                paragraph.add_run(stripped.lstrip("> "))
            elif re.match(r"^[-*]\s+", stripped):
                paragraph = doc.add_paragraph(style="List Bullet")
                paragraph.add_run(re.sub(r"^[-*]\s+", "", stripped))
            elif re.match(r"^\d+[.)]\s+", stripped):
                paragraph = doc.add_paragraph(style="List Number")
                paragraph.add_run(re.sub(r"^\d+[.)]\s+", "", stripped))
            else:
                paragraph = doc.add_paragraph(stripped)
            for run in paragraph.runs:
                _set_run_font(run, size=SMALL_SIZE if paragraph.style.name == "Long Quotation" else BODY_SIZE)


def _add_pageref(run, bookmark: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark} \\h "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "0"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, placeholder, end):
        run._r.append(element)


def _add_toc(doc, entries):
    for entry in entries:
        if entry["level"] > 3:
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm((entry["level"] - 1) * 0.6)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(16 - ((entry["level"] - 1) * 0.6)),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        _set_run_font(paragraph.add_run(entry["label"] + "\t"))
        page_run = paragraph.add_run()
        _set_run_font(page_run)
        _add_pageref(page_run, entry["bookmark"])


def _mark_update_fields(doc):
    settings = doc.settings._element
    current = settings.find(qn("w:updateFields"))
    if current is None:
        current = OxmlElement("w:updateFields")
        settings.append(current)
    current.set(qn("w:val"), "true")


def _strip_manual_numbering(title: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*[.)]?\s+", "", _safe(title)).strip()


def _add_heading(doc, label: str, level: int, *, page_break=False, bookmark=None):
    paragraph = doc.add_paragraph(style=f"Heading {min(max(level, 1), 5)}")
    paragraph.paragraph_format.page_break_before = page_break
    paragraph.paragraph_format.keep_with_next = True
    text = label.upper() if level == 1 else label
    run = paragraph.add_run(text)
    _set_run_font(run, bold=True)
    if bookmark:
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark["id"]))
        start.set(qn("w:name"), bookmark["name"])
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark["id"]))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)
    return paragraph


def _add_section_tree(doc, section, number_parts: list[int], bookmarks):
    number = ".".join(str(item) for item in number_parts)
    title = _strip_manual_numbering(section.title) or "Seção sem título"
    label_number = f"{number}." if section.level == 1 else number
    _add_heading(
        doc,
        f"{label_number} {title}",
        min(section.level, 5),
        page_break=section.level == 1,
        bookmark=bookmarks[number],
    )
    _add_prose(doc, section.content)
    for index, child in enumerate(section.children.all(), start=1):
        _add_section_tree(doc, child, [*number_parts, index], bookmarks)


def _approval_date(value) -> str:
    if not value:
        return "____/____/______"
    return value.strftime("%d/%m/%Y")


def _display_institution(value: str) -> str:
    text = _safe(value)
    if text.upper() == "SEMINÁRIO PRESBITERIANO DO NORTE - SPN":
        return "Seminário Presbiteriano do Norte – SPN"
    return text.replace(" - ", " – ")


def _display_location(value: str) -> str:
    return _safe(value).replace(" - ", " – ")


def _add_nature_block(doc, work, *, include_advisor=True):
    nature = doc.add_paragraph()
    nature.paragraph_format.left_indent = Cm(8)
    nature.paragraph_format.first_line_indent = Cm(0)
    nature.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    nature.paragraph_format.space_after = Pt(0)
    nature.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_run_font(nature.add_run(_safe(work.nature_text)), size=BODY_SIZE)
    if include_advisor and work.advisor_name:
        adviser = doc.add_paragraph()
        adviser.paragraph_format.left_indent = Cm(8)
        adviser.paragraph_format.first_line_indent = Cm(0)
        adviser.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        adviser.paragraph_format.space_before = Pt(34)
        adviser.paragraph_format.space_after = Pt(0)
        adviser.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(adviser.add_run("Orientador: "), size=BODY_SIZE, bold=True)
        _set_run_font(
            adviser.add_run(
                " ".join(filter(None, [_safe(work.advisor_title), _safe(work.advisor_name)]))
            ),
            size=BODY_SIZE,
        )


def _add_examiner(doc, title: str, name: str, institution: str, role: str, *, before=24):
    if not name:
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "000000")
    borders.append(top)
    p_pr.append(borders)
    examiner = " ".join(filter(None, [_safe(title), _safe(name)]))
    _set_run_font(paragraph.add_run(f"{examiner} ({role})"), size=SMALL_SIZE)
    if institution:
        institution_p = _add_centered(
            doc, _display_institution(institution), size=SMALL_SIZE, after=0
        )
        institution_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_front_matter(doc, work, toc_entries):
    # Capa institucional conforme o padrão visual recorrente nos trabalhos SPN de 2025.
    _add_spn_logo(doc)
    for line in [work.institution_line_1, work.institution_line_2, work.institution_line_3, work.institution_line_4, work.course_name]:
        _add_centered(
            doc,
            _safe(line).replace(" - ", " – "),
            bold=False,
            size=12,
            after=0,
            uppercase=True,
        )
    _add_spacer(doc, 67)
    _add_centered(doc, work.author_name or "NOME DO AUTOR", bold=False, uppercase=True)
    _add_spacer(doc, 129)
    title = work.title or "TÍTULO DA MONOGRAFIA"
    _add_centered(doc, title, bold=True, uppercase=True)
    if work.subtitle:
        _add_centered(doc, work.subtitle, bold=True)
    _add_spacer(doc, 160 if work.subtitle else 220)
    _add_centered(doc, _display_location(work.city), uppercase=True)
    _add_centered(doc, work.year)
    _page_break(doc)

    # Folha de rosto.
    _add_centered(doc, work.author_name or "NOME DO AUTOR", uppercase=True)
    _add_spacer(doc, 102)
    _add_centered(doc, title, bold=True, uppercase=True)
    if work.subtitle:
        _add_centered(doc, work.subtitle, bold=True)
    _add_spacer(doc, 166)
    _add_nature_block(doc, work)
    _add_spacer(doc, 129)
    _add_centered(doc, _display_location(work.city), uppercase=True)
    _add_centered(doc, work.year)
    _page_break(doc)

    # Folha de aprovação.
    _add_centered(doc, work.author_name or "NOME DO AUTOR", uppercase=True)
    _add_spacer(doc, 62 if work.subtitle else 146)
    _add_centered(doc, title, bold=True, uppercase=True)
    if work.subtitle:
        _add_centered(doc, work.subtitle, bold=True)
    _add_spacer(doc, 70 if work.subtitle else 38)
    _add_nature_block(doc, work)
    date_p = doc.add_paragraph(f"Data de aprovação: {_approval_date(work.approval_date)}")
    date_p.paragraph_format.left_indent = Cm(8)
    date_p.paragraph_format.first_line_indent = Cm(0)
    date_p.paragraph_format.space_before = Pt(20)
    date_p.paragraph_format.space_after = Pt(0)
    date_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _set_run_font(date_p.runs[0], size=BODY_SIZE)
    _add_spacer(doc, 46)
    _add_centered(doc, "BANCA EXAMINADORA", bold=True, after=0)
    _add_examiner(
        doc,
        work.advisor_title,
        work.advisor_name,
        work.institution_line_4,
        "Orientador",
        before=28,
    )
    _add_examiner(
        doc,
        work.examiner_internal_title,
        work.examiner_internal_name,
        work.examiner_internal_institution,
        "Examinador Interno",
        before=40,
    )
    _add_examiner(
        doc,
        work.examiner_external_title,
        work.examiner_external_name,
        work.examiner_external_institution,
        "Examinador Externo",
        before=40,
    )

    if work.dedication:
        _page_break(doc)
        _add_spacer(
            doc,
            _bottom_block_spacer(
                work.dedication, base=590, reduction_per_line=12
            ),
        )
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(8)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run_font(
            paragraph.add_run(_safe(work.dedication)), size=BODY_SIZE, italic=True
        )
    if work.acknowledgements:
        _page_break(doc)
        _pretext_heading(doc, "Agradecimentos")
        _add_prose(doc, work.acknowledgements)
    if work.epigraph_text:
        _page_break(doc)
        epigraph_block = "\n".join(
            filter(None, [_safe(work.epigraph_text), _safe(work.epigraph_author)])
        )
        _add_spacer(
            doc,
            _bottom_block_spacer(
                epigraph_block, base=590, reduction_per_line=18
            ),
        )
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(8)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_run_font(
            paragraph.add_run(_safe(work.epigraph_text)), size=BODY_SIZE, italic=True
        )
        if work.epigraph_author:
            author = doc.add_paragraph()
            author.paragraph_format.left_indent = Cm(8)
            author.paragraph_format.first_line_indent = Cm(0)
            author.paragraph_format.line_spacing = 1.5
            author.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_run_font(
                author.add_run(_safe(work.epigraph_author)),
                size=BODY_SIZE,
                italic=True,
            )
    if work.confessional_content or work.confessional_references:
        _page_break(doc)
        _pretext_heading(doc, "Base confessional")
        if work.confessional_title:
            _add_centered(doc, work.confessional_title, bold=True, after=4)
        if work.confessional_subtitle:
            subtitle = _add_centered(doc, work.confessional_subtitle, after=12)
            for run in subtitle.runs:
                _set_run_font(run, italic=True)
        _add_prose(doc, work.confessional_content)
        if work.confessional_references:
            _add_spacer(doc, 8)
            ref = doc.add_paragraph()
            ref.paragraph_format.first_line_indent = Cm(0)
            ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            _set_run_font(ref.add_run("Referências: "), size=10, bold=True)
            _set_run_font(ref.add_run(_safe(work.confessional_references)), size=10)

    _page_break(doc)
    _pretext_heading(doc, "Resumo")
    _add_prose(doc, work.abstract_pt or "[Resumo ainda não preenchido.]")
    keywords = doc.add_paragraph()
    keywords.paragraph_format.first_line_indent = Cm(0)
    keywords.paragraph_format.space_before = Pt(12)
    _set_run_font(keywords.add_run("Palavras-chave: "), bold=True)
    _set_run_font(keywords.add_run(_safe(work.keywords_pt) or "[não informadas]"))

    if work.abstract_en or work.keywords_en:
        _page_break(doc)
        _pretext_heading(doc, "Abstract")
        _add_prose(doc, work.abstract_en)
        keywords = doc.add_paragraph()
        keywords.paragraph_format.first_line_indent = Cm(0)
        keywords.paragraph_format.space_before = Pt(12)
        _set_run_font(keywords.add_run("Keywords: "), bold=True)
        _set_run_font(keywords.add_run(_safe(work.keywords_en)))

    if work.abbreviations:
        _page_break(doc)
        _pretext_heading(doc, "Lista de abreviaturas e siglas")
        for line in _safe(work.abbreviations).splitlines():
            if line.strip():
                paragraph = doc.add_paragraph(line.strip())
                paragraph.paragraph_format.first_line_indent = Cm(0)
    if work.symbols:
        _page_break(doc)
        _pretext_heading(doc, "Lista de símbolos")
        for line in _safe(work.symbols).splitlines():
            if line.strip():
                paragraph = doc.add_paragraph(line.strip())
                paragraph.paragraph_format.first_line_indent = Cm(0)

    _page_break(doc)
    _pretext_heading(doc, "Sumário")
    _add_toc(doc, toc_entries)


def _reference_text_sort_key(text: str):
    normalized = unicodedata.normalize("NFKD", _safe(text))
    return "".join(char for char in normalized if not unicodedata.combining(char)).casefold()


def _add_references(
    doc, publications: Iterable, imported_references: Iterable = (), *, bookmark=None
):
    rows = []
    for publication in publications:
        reference_text = format_reference(publication)
        rows.append(
            {
                "text": reference_text,
                "title": _safe(publication.title),
                "subtitle": _safe(publication.subtitle),
                "sort": _reference_text_sort_key(reference_text),
            }
        )
    rows.extend(
        {
            "text": _safe(reference.text),
            "title": "",
            "subtitle": "",
            "sort": _reference_text_sort_key(reference.text),
        }
        for reference in imported_references
    )
    unique_rows = []
    seen = set()
    for row in sorted(rows, key=lambda item: item["sort"]):
        identity = re.sub(r"\s+", " ", row["text"]).strip().casefold()
        if not identity or identity in seen:
            continue
        seen.add(identity)
        unique_rows.append(row)
    _add_heading(doc, "REFERÊNCIAS", 1, page_break=True, bookmark=bookmark)
    if not unique_rows:
        paragraph = doc.add_paragraph("[Nenhuma referência foi salva.]" )
        paragraph.paragraph_format.first_line_indent = Cm(0)
        return
    for row in unique_rows:
        text = row["text"]
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(12)
        if row["title"]:
            before, matched, after = text.partition(row["title"])
        else:
            before, matched, after = "", "", text
        _set_run_font(paragraph.add_run(before))
        if matched:
            title_text = matched
            if row["subtitle"] and row["subtitle"] in after:
                prefix, subtitle, suffix = after.partition(row["subtitle"])
                title_text += prefix + subtitle
                after = suffix
            _set_run_font(paragraph.add_run(title_text), bold=True)
        _set_run_font(paragraph.add_run(after))


def _build_navigation(work, top_sections):
    entries = []
    bookmarks = {}
    next_id = 1

    def add(key: str, label: str, level: int):
        nonlocal next_id
        bookmark = {"name": f"spn_{re.sub(r'[^A-Za-z0-9_]', '_', key)}", "id": next_id}
        next_id += 1
        bookmarks[key] = bookmark
        entries.append({"label": label, "level": level, "bookmark": bookmark["name"]})

    def walk(section, number_parts):
        number = ".".join(str(item) for item in number_parts)
        title = _strip_manual_numbering(section.title) or "Seção sem título"
        label_number = f"{number}." if section.level == 1 else number
        label = f"{label_number} {title.upper() if section.level == 1 else title}"
        add(number, label, section.level)
        for child_index, child in enumerate(section.children.all(), start=1):
            walk(child, [*number_parts, child_index])

    add("1", "1. INTRODUÇÃO", 1)
    for index, section in enumerate(top_sections, start=2):
        walk(section, [index])
    conclusion_number = 2 + len(top_sections)
    add(str(conclusion_number), f"{conclusion_number}. CONSIDERAÇÕES FINAIS", 1)
    add("references", "REFERÊNCIAS", 1)
    if work.glossary:
        add("glossary", "GLOSSÁRIO", 1)
    if work.appendices:
        add("appendices", "APÊNDICES", 1)
    if work.annexes:
        add("annexes", "ANEXOS", 1)
    return entries, bookmarks


def build_monograph_docx(work) -> BytesIO:
    """Retorna um DOCX editável. Campos automáticos são atualizados ao abrir no Word."""
    doc = Document()
    doc.core_properties.title = _safe(work.title) or "Monografia SPN"
    doc.core_properties.author = _safe(work.author_name)
    doc.core_properties.subject = "Monografia de Teologia - Seminário Presbiteriano do Norte"
    doc.core_properties.keywords = _safe(work.keywords_pt or work.planning_keywords)
    doc.core_properties.comments = "Gerado pelo aplicativo Monografia SPN em conformidade estrutural com a ABNT."
    _configure_styles(doc)
    _configure_section(doc.sections[0])
    _set_page_number_start(doc.sections[0], 0)  # A capa não entra na contagem.

    top_sections = list(
        work.sections.filter(parent__isnull=True).prefetch_related(
            "children__children__children__children"
        )
    )
    toc_entries, bookmarks = _build_navigation(work, top_sections)
    _add_front_matter(doc, work, toc_entries)

    textual_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(textual_section)
    _continue_page_numbering(textual_section)
    _add_page_number(textual_section)

    _add_heading(doc, "1. INTRODUÇÃO", 1, bookmark=bookmarks["1"])
    _add_prose(doc, work.introduction or "[Introdução ainda não preenchida.]")

    for index, section in enumerate(top_sections, start=2):
        _add_section_tree(doc, section, [index], bookmarks)

    conclusion_number = 2 + len(top_sections)
    _add_heading(
        doc,
        f"{conclusion_number}. CONSIDERAÇÕES FINAIS",
        1,
        page_break=True,
        bookmark=bookmarks[str(conclusion_number)],
    )
    _add_prose(doc, work.conclusion or "[Considerações finais ainda não preenchidas.]")
    _add_references(
        doc,
        work.publications.all(),
        work.reference_entries.all(),
        bookmark=bookmarks["references"],
    )

    if work.glossary:
        _add_heading(doc, "GLOSSÁRIO", 1, page_break=True, bookmark=bookmarks["glossary"])
        _add_prose(doc, work.glossary)
    if work.appendices:
        _add_heading(doc, "APÊNDICES", 1, page_break=True, bookmark=bookmarks["appendices"])
        _add_prose(doc, work.appendices)
    if work.annexes:
        _add_heading(doc, "ANEXOS", 1, page_break=True, bookmark=bookmarks["annexes"])
        _add_prose(doc, work.annexes)

    _mark_update_fields(doc)
    raw_output = BytesIO()
    doc.save(raw_output)
    notes_by_marker = {
        str(note.marker): note.reference_text for note in work.citation_notes.all()
    }
    return inject_footnotes(raw_output.getvalue(), notes_by_marker)
