"""Extração defensiva de listas de referências ABNT em DOCX ou PDF."""

from __future__ import annotations

from hashlib import sha256
from io import BytesIO
from pathlib import Path
import re
import unicodedata
import zipfile

from docx import Document


MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_UNCOMPRESSED_DOCX_BYTES = 40 * 1024 * 1024
MAX_PDF_PAGES = 80
MAX_EXTRACTED_TEXT_CHARS = 500_000
MAX_REFERENCES = 500
MAX_REFERENCE_CHARS = 4000
SKIPPED_HEADINGS = {
    "bibliografia",
    "bibliografia consultada",
    "referências",
    "referências bibliográficas",
}


class ReferenceImportError(ValueError):
    pass


def normalize_reference(text: str) -> str:
    text = unicodedata.normalize("NFC", str(text or ""))
    text = "".join(char for char in text if char in "\n\t" or ord(char) >= 32)
    return re.sub(r"\s+", " ", text).strip()


def reference_checksum(text: str) -> str:
    normalized = normalize_reference(text).casefold()
    return sha256(normalized.encode("utf-8")).hexdigest()


def _usable_reference(text: str) -> str | None:
    normalized = normalize_reference(text)
    if not normalized or normalized.casefold().rstrip(":") in SKIPPED_HEADINGS:
        return None
    if len(normalized) > MAX_REFERENCE_CHARS:
        raise ReferenceImportError(
            "Uma das referências ultrapassa 4.000 caracteres. Revise o arquivo e tente novamente."
        )
    return normalized


def _extract_docx(data: bytes) -> list[str]:
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            if any(".." in Path(item.filename).parts for item in archive.infolist()):
                raise ReferenceImportError("O DOCX contém caminhos internos inválidos.")
            uncompressed = sum(item.file_size for item in archive.infolist())
            if uncompressed > MAX_UNCOMPRESSED_DOCX_BYTES:
                raise ReferenceImportError("O DOCX expandido excede o limite de segurança.")
    except zipfile.BadZipFile as exc:
        raise ReferenceImportError("O arquivo DOCX está corrompido ou não é um DOCX válido.") from exc

    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise ReferenceImportError("Não foi possível ler o conteúdo do DOCX.") from exc

    candidates = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                candidates.extend(paragraph.text for paragraph in cell.paragraphs)
    return [reference for text in candidates if (reference := _usable_reference(text))]


_SURNAME_START = re.compile(
    r"^[A-ZÁÀÂÃÉÊÍÓÔÕÚÜÇ][A-ZÁÀÂÃÉÊÍÓÔÕÚÜÇ'’\- ]{1,80},\s+\S"
)
_CORPORATE_START = re.compile(
    r"^(?:BRASIL|[A-ZÁÀÂÃÉÊÍÓÔÕÚÜÇ][A-Z0-9ÁÀÂÃÉÊÍÓÔÕÚÜÇ'’()\- ]{3,80})\.\s+\S"
)


def _looks_like_reference_start(line: str) -> bool:
    return bool(_SURNAME_START.match(line) or _CORPORATE_START.match(line))


def _references_from_pdf_text(text: str) -> list[str]:
    lines = [normalize_reference(line) for line in str(text or "").splitlines()]
    references: list[str] = []
    buffer: list[str] = []

    def flush():
        if not buffer:
            return
        candidate = _usable_reference(" ".join(buffer))
        buffer.clear()
        if candidate:
            references.append(candidate)

    for line in lines:
        if not line:
            flush()
            continue
        if line.casefold().rstrip(":") in SKIPPED_HEADINGS or re.fullmatch(r"\d{1,4}", line):
            continue
        if buffer and _looks_like_reference_start(line):
            flush()
        buffer.append(line)
    flush()
    return references


def _extract_pdf(data: bytes) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ReferenceImportError(
            "O leitor de PDF ainda não está instalado no servidor. Atualize as dependências da aplicação."
        ) from exc

    try:
        reader = PdfReader(BytesIO(data), strict=False)
        if reader.is_encrypted and not reader.decrypt(""):
            raise ReferenceImportError("O PDF está protegido por senha.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ReferenceImportError(
                f"O PDF possui mais de {MAX_PDF_PAGES} páginas; envie apenas a lista de referências."
            )
        pages = []
        extracted_size = 0
        for page in reader.pages:
            page_text = page.extract_text() or ""
            extracted_size += len(page_text)
            if extracted_size > MAX_EXTRACTED_TEXT_CHARS:
                raise ReferenceImportError(
                    "O texto extraído do PDF excede o limite de segurança. Envie apenas a lista de referências."
                )
            pages.append(page_text)
        text = "\n\n".join(pages)
    except ReferenceImportError:
        raise
    except Exception as exc:
        raise ReferenceImportError("Não foi possível extrair o texto do PDF.") from exc
    if not text.strip():
        raise ReferenceImportError(
            "O PDF não contém texto selecionável. Converta a lista por OCR ou envie um DOCX."
        )
    return _references_from_pdf_text(text)


def extract_references(uploaded_file) -> list[str]:
    """Retorna referências normalizadas, preservando a ordem do arquivo."""
    filename = Path(str(getattr(uploaded_file, "name", "") or "")).name
    extension = Path(filename).suffix.casefold()
    if extension not in {".docx", ".pdf"}:
        raise ReferenceImportError("Envie um arquivo no formato DOCX ou PDF.")

    declared_size = int(getattr(uploaded_file, "size", 0) or 0)
    if declared_size > MAX_UPLOAD_BYTES:
        raise ReferenceImportError("O arquivo excede o limite de 5 MB.")
    data = uploaded_file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise ReferenceImportError("O arquivo excede o limite de 5 MB.")
    if not data:
        raise ReferenceImportError("O arquivo enviado está vazio.")

    references = _extract_docx(data) if extension == ".docx" else _extract_pdf(data)
    unique: list[str] = []
    seen: set[str] = set()
    for reference in references:
        checksum = reference_checksum(reference)
        if checksum in seen:
            continue
        seen.add(checksum)
        unique.append(reference)
        if len(unique) > MAX_REFERENCES:
            raise ReferenceImportError(
                f"O arquivo contém mais de {MAX_REFERENCES} referências. Divida a lista em arquivos menores."
            )
    if not unique:
        raise ReferenceImportError("Nenhuma referência foi encontrada no arquivo.")
    return unique
