from datetime import date
from io import BytesIO
import json
import zipfile
from unittest.mock import patch

from django.contrib.auth.models import User
from django.core import signing
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from .guidance import PARTS
from .models import (
    AIRevision,
    CitationNote,
    Monograph,
    Profile,
    Publication,
    ReferenceEntry,
    Section,
)
from .services.abnt import author_reference, citation_label, format_reference, trusted_publication_url
from .services.ai_gateway import AcademicRevision, RevisionSuggestion, _call_interactions, _extract_text
from .services.reference_import import _references_from_pdf_text, reference_checksum


class AppTestCase(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="aluno", email="aluno@example.com", password="senha-forte-123",
            first_name="João", last_name="Calvino",
        )
        Profile.objects.create(user=self.user, full_name="João Calvino", onboarding_seen=True)
        self.other = User.objects.create_user(username="outro", password="senha-forte-456")
        self.work = Monograph.objects.create(
            owner=self.user,
            author_name="João Calvino",
            title="A centralidade de Cristo na proclamação contemporânea",
            theme="Pregação cristocêntrica",
            delimitation="A pregação reformada no contexto urbano brasileiro",
            research_problem="Como a centralidade de Cristo orienta a proclamação contemporânea?",
            general_objective="Analisar a centralidade de Cristo na proclamação contemporânea.",
            specific_objectives="Examinar o contexto.\nSistematizar os fundamentos.\nAvaliar implicações pastorais.",
            justification="A pesquisa contribui para o debate homilético e para a prática pastoral.",
            methodology="Pesquisa bibliográfica, qualitativa e de orientação teológico-sistemática.",
            abstract_pt="Este estudo analisa a centralidade de Cristo na proclamação contemporânea, mediante pesquisa bibliográfica.",
            keywords_pt="Cristo; Pregação; Teologia reformada.",
            introduction="A proclamação cristã enfrenta desafios que exigem fidelidade bíblica e clareza teológica.",
            conclusion="Conclui-se que a centralidade de Cristo oferece o eixo teológico para a proclamação fiel.",
        )
        self.section = Section.objects.create(
            monograph=self.work, title="Fundamentos bíblico-teológicos", order=1, level=1,
            content="A Escritura apresenta Cristo como o centro da revelação e da proclamação cristã.",
        )
        self.client.force_login(self.user)

    def post_json(self, url, payload):
        return self.client.post(url, data=json.dumps(payload), content_type="application/json")

    def test_login_redirect_uses_namespaced_home_route(self):
        self.client.logout()
        response = self.client.post(
            reverse("login"),
            {"username": "aluno", "password": "senha-forte-123"},
        )
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse("works:home"))

    def test_all_workspace_parts_render(self):
        for slug in PARTS:
            with self.subTest(slug=slug):
                response = self.client.get(reverse("works:workspace", args=[self.work.pk, slug]))
                self.assertEqual(response.status_code, 200)
                self.assertContains(response, PARTS[slug]["label"])

    def test_credit_and_reference_controls_render(self):
        reference_text = "CALVINO, João. As institutas. São Paulo: Cultura Cristã, 2006."
        reference = ReferenceEntry.objects.create(
            monograph=self.work,
            text=reference_text,
            checksum=reference_checksum(reference_text),
            order=1,
        )
        note = CitationNote.objects.create(
            monograph=self.work,
            target_key="monograph:introduction",
            sequence=1,
            reference_text=reference_text,
            reference_entry=reference,
        )
        self.work.introduction += note.token
        self.work.save(update_fields=["introduction", "updated_at"])

        introduction = self.client.get(
            reverse("works:workspace", args=[self.work.pk, "introducao"])
        )
        self.assertContains(introduction, "Desenvolvido por: Rodrigo Niskier")
        self.assertContains(introduction, "data-citation-trigger")
        self.assertContains(introduction, "data-citation-editor")
        self.assertContains(introduction, 'id="citation-notes-data"')
        self.assertContains(introduction, "CALVINO")

        references = self.client.get(
            reverse("works:workspace", args=[self.work.pk, "referencias"])
        )
        self.assertContains(references, 'accept=".docx,.pdf,')
        self.assertContains(references, "Importar lista pronta")

    def test_dashboard_links_to_citarn(self):
        response = self.client.get(reverse("works:dashboard"))
        self.assertContains(response, reverse("references:dashboard"))
        self.assertContains(response, "CitaRN")

    def test_new_monograph_has_spn_structure_starters(self):
        response = self.client.post(reverse("works:create_monograph"))
        created = Monograph.objects.filter(owner=self.user).latest("created_at")
        self.assertRedirects(response, reverse("works:workspace", args=[created.pk, "planejamento"]))
        self.assertEqual(created.author_name, "João Calvino")
        self.assertEqual(created.sections.filter(level=1).count(), 3)
        self.assertTrue(created.sections.filter(title__icontains="confessionais").exists())

    def test_autosave_is_allowlisted_and_scoped(self):
        url = reverse("works:autosave", args=[self.work.pk])
        response = self.post_json(url, {"field": "theme", "value": "Novo tema delimitado"})
        self.assertEqual(response.status_code, 200)
        self.work.refresh_from_db()
        self.assertEqual(self.work.theme, "Novo tema delimitado")
        forbidden = self.post_json(url, {"field": "owner", "value": self.other.pk})
        self.assertEqual(forbidden.status_code, 403)
        self.client.force_login(self.other)
        hidden = self.post_json(url, {"field": "theme", "value": "Tentativa"})
        self.assertEqual(hidden.status_code, 404)

    def test_sections_support_hierarchy_and_protect_ownership(self):
        add_url = reverse("works:add_section", args=[self.work.pk])
        response = self.post_json(add_url, {"parent_id": self.section.pk, "title": "Exegese do texto"})
        self.assertEqual(response.status_code, 200)
        child = Section.objects.get(pk=response.json()["section"]["id"])
        self.assertEqual(child.level, 2)
        update_url = reverse("works:update_section", args=[self.work.pk, child.pk])
        updated = self.post_json(update_url, {"field": "content", "value": "Análise exegética desenvolvida."})
        self.assertEqual(updated.status_code, 200)
        child.refresh_from_db()
        self.assertEqual(child.content, "Análise exegética desenvolvida.")
        rendered = self.client.get(reverse("works:workspace", args=[self.work.pk, "desenvolvimento"]))
        self.assertContains(rendered, 'class="section-number">1.1</span>')

    @override_settings(GEMINI_API_KEY="test-key")
    @patch("works.views.revise_text")
    def test_ai_revision_requires_explicit_acceptance_and_rejects_stale_text(self, revise):
        revise.return_value = (
            AcademicRevision(
                revised_text="A Escritura apresenta Cristo como centro da revelação e da proclamação cristã.",
                summary="Ajuste de concisão sem mudança da tese.",
                suggestions=[RevisionSuggestion(category="clareza", original_excerpt="o centro", proposed_change="centro", reason="Evita artigo dispensável.", priority="baixa")],
                warnings=[], voice_preserved=True, confidence=.91,
            ),
            "gemini-test",
        )
        review_url = reverse("works:ai_review", args=[self.work.pk])
        response = self.post_json(review_url, {"target_type": "section", "section_id": self.section.pk, "action": "review"})
        self.assertEqual(response.status_code, 200)
        revision = AIRevision.objects.get(pk=response.json()["revision"]["id"])
        self.section.refresh_from_db()
        self.assertNotEqual(self.section.content, revision.proposed_text)

        self.section.content += " Alteração posterior."
        self.section.save()
        stale = self.post_json(reverse("works:accept_revision", args=[self.work.pk, revision.pk]), {})
        self.assertEqual(stale.status_code, 409)
        self.assertFalse(AIRevision.objects.get(pk=revision.pk).accepted)

    @patch("works.views.search_publications")
    def test_research_results_are_signed_before_save(self, search):
        search.return_value = ([{
            "title": "Institutas da Religião Cristã", "subtitle": "", "authors": ["João Calvino"],
            "year": 2006, "source_type": "book", "url": "https://openlibrary.org/works/OL123W",
            "provider": "Open Library", "city": "São Paulo", "publisher": "Cultura Cristã", "edition": "",
            "container_title": "", "volume": "", "issue": "", "pages": "", "doi": "", "isbn": "123",
            "language": "por", "relevance": .92,
        }], [])
        result = self.post_json(reverse("works:research_search", args=[self.work.pk]), {"query": "teologia reformada", "mode": "classic"})
        self.assertEqual(result.status_code, 200)
        item = result.json()["results"][0]
        saved = self.post_json(reverse("works:save_publication", args=[self.work.pk]), {"token": item["token"]})
        self.assertEqual(saved.status_code, 200)
        publication = Publication.objects.get(monograph=self.work)
        self.assertEqual(publication.provider, "Open Library")
        tampered = item["token"] + "x"
        rejected = self.post_json(reverse("works:save_publication", args=[self.work.pk]), {"token": tampered})
        self.assertEqual(rejected.status_code, 403)

    def test_docx_export_contains_spn_and_abnt_structure(self):
        Publication.objects.create(
            monograph=self.work, source_type="book", title="A missão da igreja", authors=["Michael Horton"],
            year=2012, city="São Paulo", publisher="Cultura Cristã", url="https://openlibrary.org/works/OL1W",
            provider="Open Library", access_date=date(2026, 8, 13),
        )
        response = self.client.get(reverse("works:export_docx", args=[self.work.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        doc = Document(BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("SEMINÁRIO PRESBITERIANO DO NORTE – SPN", text)
        self.assertIn("1. INTRODUÇÃO", text)
        self.assertIn("2. FUNDAMENTOS BÍBLICO-TEOLÓGICOS", text)
        self.assertIn("3. CONSIDERAÇÕES FINAIS", text)
        self.assertIn("BANCA EXAMINADORA", text)
        self.assertIn("REFERÊNCIAS", text)
        toc_heading_index = next(
            index
            for index, paragraph in enumerate(doc.paragraphs)
            if paragraph.text == "SUMÁRIO"
        )
        first_toc_entry = doc.paragraphs[toc_heading_index + 1]
        self.assertEqual(first_toc_entry.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(
            first_toc_entry.paragraph_format.line_spacing_rule,
            WD_LINE_SPACING.SINGLE,
        )
        self.assertAlmostEqual(first_toc_entry.runs[0].font.size.pt, 10, places=1)
        self.assertTrue(first_toc_entry.runs[0].bold)
        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            self.assertTrue(
                any(name.startswith("word/media/") for name in archive.namelist())
            )
        first = doc.sections[0]
        self.assertAlmostEqual(first.top_margin.cm, 3, places=1)
        self.assertAlmostEqual(first.left_margin.cm, 3, places=1)
        self.assertAlmostEqual(first.right_margin.cm, 2, places=1)

    def test_docx_reference_list_import_and_deduplication(self):
        source = Document()
        source.add_paragraph("REFERÊNCIAS")
        source.add_paragraph(
            "CALVINO, João. As institutas. São Paulo: Cultura Cristã, 2006."
        )
        source.add_paragraph(
            "HORTON, Michael. A missão da igreja. São Paulo: Cultura Cristã, 2012."
        )
        content = BytesIO()
        source.save(content)
        uploaded = SimpleUploadedFile(
            "referencias.docx",
            content.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = self.client.post(
            reverse("works:import_references", args=[self.work.pk]),
            {"reference_file": uploaded},
        )
        self.assertRedirects(
            response,
            reverse("works:workspace", args=[self.work.pk, "referencias"]),
        )
        self.assertEqual(self.work.reference_entries.count(), 2)

        content.seek(0)
        duplicate_upload = SimpleUploadedFile("referencias.docx", content.getvalue())
        self.client.post(
            reverse("works:import_references", args=[self.work.pk]),
            {"reference_file": duplicate_upload},
        )
        self.assertEqual(self.work.reference_entries.count(), 2)

    def test_citation_note_is_inserted_at_cursor_and_can_be_removed(self):
        reference_text = (
            "CALVINO, João. As institutas. São Paulo: Cultura Cristã, 2006."
        )
        reference = ReferenceEntry.objects.create(
            monograph=self.work,
            text=reference_text,
            source_filename="referencias.docx",
            checksum=reference_checksum(reference_text),
            order=1,
        )
        before = "A proclamação cristã"
        response = self.post_json(
            reverse("works:create_citation_note", args=[self.work.pk]),
            {
                "target_key": "monograph:introduction",
                "current_text": self.work.introduction,
                "before_text": before,
                "source_kind": "imported",
                "source_id": reference.pk,
                "locator": "p. 42.",
            },
        )
        self.assertEqual(response.status_code, 200)
        note = CitationNote.objects.get(monograph=self.work)
        self.work.refresh_from_db()
        self.assertEqual(note.sequence, 1)
        self.assertIn(note.token, self.work.introduction)
        self.assertTrue(self.work.introduction.startswith(before + note.token))
        self.assertTrue(note.reference_text.endswith("p. 42."))

        removed = self.post_json(
            reverse("works:delete_citation_note", args=[self.work.pk, note.pk]), {}
        )
        self.assertEqual(removed.status_code, 200)
        self.work.refresh_from_db()
        self.assertNotIn(note.token, self.work.introduction)
        self.assertFalse(CitationNote.objects.filter(pk=note.pk).exists())

    def test_docx_export_contains_true_footnote_and_imported_bibliography(self):
        reference_text = (
            "CALVINO, João. As institutas. São Paulo: Cultura Cristã, 2006."
        )
        reference = ReferenceEntry.objects.create(
            monograph=self.work,
            text=reference_text,
            source_filename="referencias.docx",
            checksum=reference_checksum(reference_text),
            order=1,
        )
        publication = Publication.objects.create(
            monograph=self.work,
            source_type="book",
            title="A missão da igreja no mundo contemporâneo",
            authors=["Michael Horton"],
            year=2012,
            city="São Paulo",
            publisher="Cultura Cristã",
            url="https://openlibrary.org/works/OL123W",
            provider="Open Library",
        )
        note = CitationNote.objects.create(
            monograph=self.work,
            target_key="monograph:introduction",
            sequence=1,
            reference_text=reference_text + " p. 42.",
            reference_entry=reference,
        )
        self.work.introduction += note.token
        self.work.save(update_fields=["introduction", "updated_at"])

        response = self.client.get(
            reverse("works:export_docx", args=[self.work.pk])
        )
        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            self.assertIn("word/footnotes.xml", archive.namelist())
            self.assertIn("word/_rels/footnotes.xml.rels", archive.namelist())
            document_xml = archive.read("word/document.xml").decode("utf-8")
            footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
            self.assertIn("footnoteReference", document_xml)
            self.assertNotIn("[[FN:", document_xml)
            self.assertIn('w:type="separator"', footnotes_xml)
            self.assertIn('w:type="continuationSeparator"', footnotes_xml)
            self.assertIn("As institutas", footnotes_xml)
            self.assertIn("p. 42.", footnotes_xml)
        exported = Document(BytesIO(response.content))
        body_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        self.assertIn(reference_text, body_text)
        horton_reference = format_reference(publication)
        self.assertLess(body_text.index(reference_text), body_text.index(horton_reference))

    def test_citation_note_cannot_use_another_users_reference(self):
        other_work = Monograph.objects.create(owner=self.other)
        foreign_text = "AUTOR, Outro. Obra alheia. Recife: Editora, 2020."
        foreign_reference = ReferenceEntry.objects.create(
            monograph=other_work,
            text=foreign_text,
            checksum=reference_checksum(foreign_text),
            order=1,
        )
        response = self.post_json(
            reverse("works:create_citation_note", args=[self.work.pk]),
            {
                "target_key": "monograph:introduction",
                "current_text": self.work.introduction,
                "before_text": self.work.introduction,
                "source_kind": "imported",
                "source_id": foreign_reference.pk,
            },
        )
        self.assertEqual(response.status_code, 404)


class ReferenceImportParsingTests(TestCase):
    def test_pdf_lines_are_grouped_into_complete_references(self):
        text = """
REFERÊNCIAS

CALVINO, João. As institutas da religião cristã.
São Paulo: Cultura Cristã, 2006.

HORTON, Michael. A missão da igreja.
São Paulo: Cultura Cristã, 2012.
"""
        references = _references_from_pdf_text(text)
        self.assertEqual(len(references), 2)
        self.assertIn("São Paulo: Cultura Cristã, 2006.", references[0])


class ABNTFormattingTests(TestCase):
    def test_reference_and_author_date_follow_current_conventions(self):
        publication = SimplePublication()
        reference = format_reference(publication)
        self.assertTrue(reference.startswith("SILVA, João."))
        self.assertIn("Disponível em: https://doi.org/10.1000/teste.", reference)
        self.assertEqual(citation_label(publication), "(Silva, 2025)")

    def test_only_https_known_catalog_links_are_accepted(self):
        self.assertTrue(trusted_publication_url("https://doi.org/10.1000/teste"))
        self.assertTrue(trusted_publication_url("https://books.google.com/books?id=123"))
        self.assertTrue(trusted_publication_url("https://www.scielo.br/j/revista/a/abc"))
        self.assertFalse(trusted_publication_url("http://doi.org/10.1000/teste"))
        self.assertFalse(trusted_publication_url("https://doi.org.evil.example/10.1000/teste"))
        self.assertFalse(trusted_publication_url("https://example.invalid/obra"))

    def test_portuguese_surname_particles_stay_with_given_names(self):
        self.assertEqual(author_reference("Maria da Silva"), "SILVA, Maria da")


class AIGatewayCompatibilityTests(TestCase):
    def test_interactions_request_uses_current_structured_output_shape(self):
        client = MockInteractionsClient()
        _call_interactions(client, "gemini-3.6-flash", "Texto de teste")
        body = client.interactions.body
        self.assertEqual(body["response_format"]["type"], "text")
        self.assertEqual(body["response_format"]["mime_type"], "application/json")
        self.assertIn("properties", body["response_format"]["schema"])
        self.assertNotIn("temperature", body["generation_config"])
        self.assertFalse(body["store"])

    def test_interactions_output_text_is_extracted(self):
        response = type("InteractionResponse", (), {"output_text": '{"ok": true}'})()
        self.assertEqual(_extract_text(response), '{"ok": true}')


class SimplePublication:
    authors = ["João Silva"]
    title = "Teologia e missão"
    subtitle = "um estudo"
    year = 2025
    source_type = "article"
    container_title = "Revista Teológica"
    volume = "10"
    issue = "2"
    pages = "10-30"
    doi = "10.1000/teste"
    url = "https://doi.org/10.1000/teste"
    access_date = date(2026, 8, 13)
    edition = ""
    city = ""
    publisher = ""


class MockInteractions:
    def __init__(self):
        self.body = None

    def create(self, **kwargs):
        self.body = kwargs
        return object()


class MockInteractionsClient:
    def __init__(self):
        self.interactions = MockInteractions()
